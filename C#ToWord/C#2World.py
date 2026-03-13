import os
import sys
from docx import Document
from docx.shared import Pt

def gather_cs_lines(root_folder):
    """递归读取 .cs 文件，并返回所有代码行（含文件头，过滤空行）"""
    cs_lines = []
    root_folder = os.path.abspath(root_folder)
    for dirpath, dirnames, filenames in os.walk(root_folder):
        filenames = sorted(filenames)
        for fn in filenames:
            if fn.lower().endswith(".cs"):
                full = os.path.join(dirpath, fn)
                rel = os.path.relpath(full, root_folder)
                cs_lines.append(f"// File: {rel}")
                try:
                    with open(full, "r", encoding="utf-8") as f:
                        for line in f:
                            line = line.rstrip("\n").rstrip("\r")
                            if not line.strip():  # ✅ 跳过空行
                                continue
                            cs_lines.append(line)
                except UnicodeDecodeError:
                    with open(full, "r", encoding="latin-1", errors="ignore") as f:
                        for line in f:
                            line = line.rstrip("\n").rstrip("\r")
                            if not line.strip():  # ✅ 跳过空行
                                continue
                            cs_lines.append(line)
    return cs_lines

def write_word_from_lines(lines, out_path, pages=60, lines_per_page=50):
    """把代码行写入 Word 文档，每页 50 行，总共 60 页"""
    total_needed = pages * lines_per_page
    if len(lines) < total_needed:
        lines += [""] * (total_needed - len(lines))
    else:
        lines = lines[:total_needed]

    doc = Document()
    for i in range(pages):
        for j in range(lines_per_page):
            idx = i * lines_per_page + j
            p = doc.add_paragraph()
            run = p.add_run(lines[idx])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        if i != pages - 1:
            doc.add_page_break()
    doc.save(out_path)
    print(f"✅ 已生成 Word 文件：{out_path}")

if __name__ == "__main__":
    print("=== C#代码生成Word文档工具（已过滤空行） ===")
    src_folder = input("请输入C#代码所在文件夹路径：").strip('" ')
    out_path = input("请输入要保存的Word文件路径（如 D:\\output\\code.docx）：").strip('" ')

    if not os.path.isdir(src_folder):
        print("❌ 输入的文件夹路径不存在！")
        sys.exit(1)

    print("正在收集代码行...")
    lines = gather_cs_lines(src_folder)
    print(f"共收集到 {len(lines)} 行代码（已过滤空行）")

    print("正在生成Word文档（共60页，每页50行）...")
    write_word_from_lines(lines, out_path)